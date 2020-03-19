from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException, StaleElementReferenceException, TimeoutException
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import pathlib
import win32com
import time
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import os
import shutil
from docx.shared import Inches, Pt
import glob
import re
import win32com.client as win32
from win32com.client import constants
import docx, docxpy
import unittest

script_path = pathlib.Path(__file__).parent.absolute()
download_path = str(script_path) + "\\Downloads"
<<<<<<< HEAD
user_login = ""
user_password = ""
=======
user_login = "akulig"
user_password = "5lazzjdzZ!!"
>>>>>>> 0bf054bdc9b8004ac889905147e468b2a2a0f664


class JiraTestsDownload(unittest.TestCase):

    @classmethod
    def setUpClass(self):  # setUpClass runs once for ALL tests
        options = webdriver.ChromeOptions()
        preferences = {"download.default_directory": download_path, "download.prompt_for_download": "false",
                       "safebrowsing.enabled": "false", 'profile.default_content_setting_values.automatic_downloads': 1}
        options.add_experimental_option("prefs", preferences)
        self.driver = webdriver.Chrome(options=options)

    def test_1_login_to_app(self):
        driver = self.driver
        driver.implicitly_wait(10)
        # user_login = input("Username: ")
        # user_password = input("Password: ")
        driver.get("https://jira.softserveinc.com/projects/SSMWE?selectedItem=com.thed.zephyr.je:zephyr-tests-page")
        login_element = driver.find_element_by_id("login-form-username")
        password_element = driver.find_element_by_id("login-form-password")
        login_element.send_keys(user_login, Keys.TAB)
        password_element.click()
        password_element.send_keys(user_password, Keys.ENTER)

    def test_2_order_tests(self):
        driver = self.driver
        wait = WebDriverWait(driver, 10)
        wait.until(EC.element_to_be_clickable((By.LINK_TEXT, "Unscheduled")))
        driver.find_element_by_link_text("Unscheduled").click()
        wait.until(EC.element_to_be_clickable((By.CLASS_NAME, "order-options")))
        driver.find_element_by_class_name("order-options").click()
        wait.until(EC.element_to_be_clickable((By.CLASS_NAME, "check-list-field-container")))
        search = driver.find_element_by_id("order-by-options-input")
        search.send_keys("Updated")
        wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='updated-1']/label"))) # wait.until(EC.element_to_be_clickable((By.XPATH,"//*[@title='Updated']")))
        search.send_keys(Keys.ENTER)
        time.sleep(3)

    def test_3_download_tests(self):
        driver = self.driver
        wait = WebDriverWait(driver, 2)
        while True:
            tests_list = driver.find_elements_by_class_name("splitview-issue-link")
            for test_element in tests_list:
                if tests_list.index(test_element) == 3:  # FOR TESTIG ONLY
                    break  # FOR TESTIG ONLY
                test_element.click()
                for x in range(0, 10):
                    try:
                        test_element_key = test_element.find_element_by_xpath(".//span[@class='issue-link-key']").text
                        wait.until(EC.presence_of_element_located((By.XPATH, "//a[@id='key-val']".format(test_element_key))))
                        # JiraTestsDownload.add_test_to_epic(self, test_element_key)
                        break
                    except (StaleElementReferenceException, TimeoutException):
                        time.sleep(0.2)
                        continue
                for x in range(0, 10):
                    try:
                        wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='viewissue-export']")))
                        driver.find_element_by_xpath("//*[@id='viewissue-export']").click()
                        wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='jira.issueviews:issue-word']")))
                        driver.find_element_by_xpath("//*[@id='jira.issueviews:issue-word']").click()
                        break
                    except (StaleElementReferenceException, TimeoutException):
                        time.sleep(0.2)
                        continue
            try:
                break  # FOR TESTIG ONLY
                wait.until(EC.element_to_be_clickable((By.XPATH, "//a[@class='icon icon-next']")))
                driver.find_element_by_xpath("//a[@class='icon icon-next']").click()
                time.sleep(3)
            except NoSuchElementException:
                break
        print("Downloading done")

    # def add_test_to_epic(self, test_element_key):
    #     driver = self.driver
    #     wait = WebDriverWait(driver, 1)
    #     driver.implicitly_wait(0)
    #     epic_name = ""
    #     try:
    #         try:
    #             wait.until(EC.presence_of_element_located((By.XPATH, "//div[@data-fieldtype='gh-epic-link']//a")))
    #             epic_name = driver.find_element_by_xpath("//div[@data-fieldtype='gh-epic-link']//a").text
    #         except (NoSuchElementException, StaleElementReferenceException, TimeoutException):
    #             wait.until(EC.presence_of_element_located((By.XPATH, "(//a[@class='lozenge'])[2]")))
    #             epic_name = driver.find_element_by_xpath("(//a[@class='lozenge'])[2]").text
    #     except (StaleElementReferenceException, TimeoutException, NoSuchElementException):
    #         pass
    #     if epic_name.strip():
    #         epics_dictionary.setdefault(epic_name, []).append(test_element_key)
    #         print(epics_dictionary)

    @classmethod
    def tearDownClass(self):
        time.sleep(1)
        self.driver.close()

#
# def rename_files():
#     os.chdir(download_path)
#     for filename in os.listdir(download_path):
#         if filename.startswith("SSMWE"):
#             os.rename(filename, filename.replace("SSMWE-", "Copy-"))


def create_dir_hierarchy():
    if not os.path.exists(download_path + "\\DocFiles"):
        os.makedirs(download_path + "\\DocFiles")
    if not os.path.exists(download_path + "\\DocxFiles"):
        os.makedirs(download_path + "\\DocxFiles")
    if not os.path.exists(download_path + "\\TestTemplates"):
        os.makedirs(download_path + "\\TestTemplates")
    if not os.path.exists(download_path + "\\Directories"):
        os.makedirs(download_path + "\\Directories")


def move_doc_files():
    # Create list of paths to .doc files
    for filename in os.listdir(download_path):
        if filename.endswith(".doc") and filename.startswith("SSMWE"):
            if os.path.exists(download_path + "\\DocFiles\\" + filename):
                os.remove(download_path + "\\DocFiles\\" + filename)
            shutil.move(download_path + "\\" + filename, download_path + "\\DocFiles")


# this function isn't standalone, use save_to_docx instead
def save_as_docx(name):  # Github conversion solution
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(download_path + "\\DocFiles\\" + name)
    doc.Activate()

    # Rename path with .docx
    new_file_abs = os.path.abspath(download_path + "\\DocxFiles\\" + name)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(new_file_abs, FileFormat=constants.wdFormatXMLDocument)
    doc.Close(False)


def save_to_docx():  # chceck obs if it doesnt work
    file_names = os.listdir(download_path + "\\DocFiles")
    for name in file_names:
        if os.path.exists((download_path + "\\DocxFiles\\" + name)[:-3] + "docx"):
            os.remove((download_path + "\\DocxFiles\\" + name)[:-3] + "docx")
        save_as_docx(name)


def read_docx_files():
    files = glob.glob(download_path + "\\DocxFiles\\SSMWE*")
    os.chdir(download_path)
    for file in files:
        if os.path.exists(download_path + "\\TestTemplates\\" + os.path.basename(file)):
            os.remove(download_path + "\\TestTemplates\\" + os.path.basename(file))
        docx_handler = docx.Document(file)
        docx_tables = docx_handler.tables
        docx_hyperlink_handler = docxpy.DOCReader(file)
        docx_hyperlink_handler.process()
        hyperlinks = docx_hyperlink_handler.data['links']
        test_scenario_hyperlink_text = str(hyperlinks[0][0])[2:-1]
        jira_id = docx_tables[0].rows[0].cells[0].text
        jira_test_id = jira_id.split()[0]

        if "Zephyr" in docx_tables[2].rows[2].cells[0].text:
            zephyr_tests = docx_tables[2].rows[2].cells[1]
        elif "Zephyr" in docx_tables[2].rows[3].cells[0].text:
            zephyr_tests = docx_tables[2].rows[3].cells[1]
        else:
            break
        zephyr_tests_table = zephyr_tests.tables

        zephyr_rows = zephyr_tests_table[0].rows  # get row id's
        zephyr_rows = zephyr_rows[1:]  # remove first cell from all rows (e.g. "Test Step", "Test Data", etc.)

        #  Test Steps
        list_of_test_steps = []
        for row in zephyr_rows:
            list_of_test_steps.append(row.cells[1].text)

        #  Test Conditions
        list_of_test_conditions = []
        for row in zephyr_rows:
            list_of_test_conditions.append(row.cells[2].text)

        #  Expected results
        list_of_exptected_results = []
        for row in zephyr_rows:
            list_of_exptected_results.append(row.cells[3].text)

        number_of_teststeps = zephyr_tests_table[0].rows[-1].cells[0].text
        file_save_path = download_path + "\\TestTemplates\\" + os.path.basename(file)

        rws_template = docx.Document(download_path + "\\SampleTestScripts1.docx")
        rws_table = rws_template.tables
        font = rws_template.styles['Normal'].font
        font.name = 'Calibri'
        paragraph = rws_template.styles['Normal'].paragraph_format
        paragraph.space_after = Pt(3)
        paragraph.left_indent = Pt(0)
        for x in range(1, int(number_of_teststeps)):
            rws_table[0].add_row()

        os.chdir(download_path)
        test_id = rws_table[0].rows[0].cells[2].paragraphs[0]
        test_id.add_run(jira_test_id)
        test_scenario = rws_table[0].rows[1].cells[2].paragraphs[0]
        test_scenario.add_run(test_scenario_hyperlink_text)
        user_story = rws_table[0].rows[4].cells[2].paragraphs[0]
        user_story.add_run(jira_test_id)

        steps_only_table = rws_table[0].rows[6:]

        # For some reason indentation formatting by using styles doesnt work for already existing table paragraphs
        # You should only use styles for newly-created elements
        # steps_only_table[0].cells[0].paragraphs[0].style = rws_template.styles['Normal']
        # So I had to do it manually
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[0].paragraphs[0].paragraph_format.left_indent = Pt(12)
            # steps_only_table[x].cells[0].paragraphs[0].runs[0].font.size = Pt(15)

        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[0].paragraphs[0].add_run(str(x + 1) + ".")
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[1].paragraphs[0].add_run(list_of_test_steps[x])
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[2].paragraphs[0].add_run(list_of_test_conditions[x])
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[3].paragraphs[0].add_run(list_of_exptected_results[x])

        rws_template.add_page_break()
        rws_template.save(file_save_path)


def move_files_to_epics():
    files = glob.glob(download_path + "\\DocxFiles\\SSMWE*")
    os.chdir(download_path)
    for file_path in files:
        docx_handler = docx.Document(file_path)
        docx_tables = docx_handler.tables
        docx_hyperlink_handler = docxpy.DOCReader(file_path)
        docx_hyperlink_handler.process()
        hyperlinks = docx_hyperlink_handler.data['links']

        epic_name = ""
        list_of_elements = []
        for x in range(len(docx_tables[2].rows)):
            list_of_elements.append(docx_tables[2].rows[x].cells[0].text)

        for index, elem in enumerate(list_of_elements):
            if "Epic/Theme" in elem:
                epic_name = docx_tables[2].rows[index].cells[1].text
                break
        for elem in list_of_elements:
            if "Epic Link" in elem:
                epic_name = (str(hyperlinks[len(hyperlinks)-1][0])[2:-1]).strip()
                for char in r'/<\>?:|"*':
                    epic_name = epic_name.replace(char, "")
                break
        if not epic_name:
            break

        os.chdir(download_path + "\\TestTemplates")
        if not os.path.exists(download_path + "\\Directories\\" + epic_name):
            os.makedirs(download_path + "\\Directories\\" + epic_name)
        if os.path.exists(download_path + "\\TestTemplates\\" + os.path.basename(file_path)):
            if os.path.exists(download_path + "\\Directories\\" + epic_name + "\\" + os.path.basename(file_path)):
                os.remove(download_path + "\\Directories\\" + epic_name + "\\" + os.path.basename(file_path))
            shutil.move(download_path + "\\TestTemplates\\" + os.path.basename(file_path), download_path + "\\Directories\\" + epic_name)


def merge_files_in_epics():
    epics_dirs = glob.glob(download_path + "\\Directories\\*")
    for epic_dir in epics_dirs:
        if not os.listdir(epic_dir) or not os.path.isdir(epic_dir):
            break
        os.chdir(epic_dir)
        files = glob.glob(epic_dir + "\\SSMWE*")
        epic_docx = docx.Document(str(script_path) + "\\EpicTemplate.docx")
        epic_docx.add_page_break()

        font = epic_docx.styles['Normal'].font
        font.name = 'Calibri'
        paragraph = epic_docx.styles['Normal'].paragraph_format
        paragraph.space_after = Pt(3)
        paragraph.left_indent = Pt(0)

        for index, file_path in enumerate(files):
            docx_handler = docx.Document(file_path)

            if index < len(files)-1:
                docx_handler.add_page_break()
            for index, element in enumerate(docx_handler.element.body):
                epic_docx.element.body.append(element)  # CHECK WHAT DOES IT DO
                if index == 4:
                    break
        epic_docx.save(epic_dir + "\\" + os.path.basename(epic_dir) + ".docx")
        time.sleep(2)
        word = win32com.client.DispatchEx("Word.Application")
        doc = word.Documents.Open(epic_dir + "\\" + os.path.basename(epic_dir) + ".docx")
        doc.TablesOfContents(1).Update()
        doc.Close(SaveChanges=True)
        word.Quit()


if __name__ == "__main__":

    # unittest.main(exit=False)  # Dont stop the program after test execution (it would skip below functions)
    # create_dir_hierarchy()
    # move_doc_files()
    # save_to_docx()
    # read_docx_files()
    # move_files_to_epics()
    merge_files_in_epics()

