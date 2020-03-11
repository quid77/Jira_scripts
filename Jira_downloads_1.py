from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import getpass
import time
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import os
import shutil
import difflib
from glob import glob
import re
import win32com.client as win32
from win32com.client import constants
from docx import Document
import docx
import glob
import fnmatch

driver = webdriver.Chrome()
driver.maximize_window()
log = input("Podaj username: ")
haslo =input("Podaj Haslo:")
wait = WebDriverWait(driver, 30)
driver.get("https://jira.softserveinc.com/projects/SSMWE?selectedItem=com.thed.zephyr.je:zephyr-tests-page")
login = driver.find_element_by_id("login-form-username")
password = driver.find_element_by_id("login-form-password")

login.send_keys(log, Keys.TAB)
password.click()
password.send_keys(haslo, Keys.ENTER)
wait.until(EC.element_to_be_clickable((By.LINK_TEXT,"Unscheduled")))
driver.find_element_by_link_text("Unscheduled").click()
wait.until(EC.element_to_be_clickable((By.CLASS_NAME,"order-options")))
driver.find_element_by_class_name("order-options").click()
wait.until(EC.element_to_be_clickable((By.CLASS_NAME,"check-list-field-container")))
search = driver.find_element_by_id("order-by-options-input")
search.send_keys("Updated")
wait.until(EC.element_to_be_clickable((By.XPATH,"//*[@id='updated-1']/label")))
search.send_keys(Keys.ENTER)
time.sleep(2)
wait.until(EC.element_to_be_clickable((By.CLASS_NAME,"issue-link-key")))
listTestow=driver.find_elements_by_class_name("issue-link-key")
ilosc_plikow = print(len(listTestow))
for test in listTestow:
     test.click()
     time.sleep(2)
     wait.until(EC.element_to_be_clickable((By.XPATH,"//*[@id='viewissue-export']")))
     driver.find_element_by_xpath("//*[@id='viewissue-export']").click()
     wait.until(EC.element_to_be_clickable((By.XPATH,"//*[@id='jira.issueviews:issue-word']")))
     driver.find_element_by_xpath("//*[@id='jira.issueviews:issue-word']").click()
time.sleep(3)
print("Pobieranie zakonczone ")

path = "C:\\Users\\jpiet\\Downloads\\"
folders = ["DocFiles"]
for el in folders:
    if not os.path.exists(path + el):
        os.makedirs(path + el)

test = os.listdir('C:\\Users\\jpiet\\Downloads')
os.chdir("C:\\Users\\jpiet\\Downloads")
for file in os.listdir("C:\\Users\\jpiet\\Downloads"):
    if file.startswith("SSMWE"):
        os.rename(file,file.replace("SSMWE-", "Copy-"))

for file in os.listdir("C:\\Users\\jpiet\\Downloads"):
    if file.endswith(".doc"):
        print(os.path.join("C:\\Users\\jpiet\\Downloads", file))
        dupa = os.path.join("C:\\Users\\jpiet\\Downloads", file)
        shutil.move(dupa, "C:\\Users\\jpiet\\Downloads\\DocFiles")
        #os.rename
# Create list of paths to .doc files
paths = os.listdir(r"C:\Users\jpiet\Downloads\DocFiles") 


def save_as_docx(path):
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate ()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)
    print(new_file_abs)
    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)

print(paths)
for path in paths:
    save_as_docx(os.path.join(r"C:\Users\jpiet\Downloads\DocFiles", path))

os.chdir("C:\\Users\\jpiet\\Downloads\\DocFiles")
doc = os.listdir("C:\\Users\\jpiet\\Downloads\\DocFiles")
for item in doc:
    if item.endswith(".doc"):
        os.remove(item)
time.sleep(10)
#Read word document
path = 'C:\\Users\\jpiet\\Downloads\\DocFiles\\Copy*'
files = glob.glob(path)
os.chdir("C:\\Users\\jpiet\\Downloads")
for file in files:

    doc = docx.Document(file)
    table = doc.tables
    titlle = table[0].rows[0].cells[0].text
    print(titlle)
    title_split = titlle.split()
    print(title_split)
    idTest= title_split[0]
    print(idTest)
    step2= table[2].rows[2].cells[0].text
    print(step2)
    if step2 == "Zephyr Teststep:":
        Testy_zefirowe = table[2].rows[2].cells[1]
    else:
        Testy_zefirowe = table[2].rows[3].cells[1]
    tabela_zefir = Testy_zefirowe.tables
    step1 = tabela_zefir[0].rows[-1].cells[0].text

    print(step1)

    tab_zef = tabela_zefir[0].rows
    tab_zef = tab_zef[1:] # obcinanie tytulu test stepow

    lista_stepow = []
    for row in tab_zef:
        lista_stepow.append(row.cells[1].text)

    lista_stepow[0]
    #Test Condition
    lista_TestCondition=[]

    for row in tab_zef:
        lista_TestCondition.append(row.cells[2].text)

    lista_ExpResult=[]
    for row in tab_zef:
        lista_ExpResult.append(row.cells[3].text)
    if step1 == "2":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts.docx")
        table1 = doc1.tables
    elif step1 == "3":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts3.docx")
        table1 = doc1.tables
    elif step1 == "4":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts4.docx")
        table1 = doc1.tables
    elif step1 == "5":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts5.docx")
        table1 = doc1.tables
    elif step1 == "6":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts6.docx")
        table1 = doc1.tables
    elif step1 == "7":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts7.docx")
        table1 = doc1.tables
    elif step1 == "8":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts8.docx")
        table1 = doc1.tables
    elif step1 == "9":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts9.docx")
        table1 = doc1.tables
    elif step1 == "10":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts10.docx")
        table1 = doc1.tables
    elif step1 == "11":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts11.docx")
        table1 = doc1.tables
    elif step1 == "12":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts12.docx")
        table1 = doc1.tables
    elif step1 == "13":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts13.docx")
        table1 = doc1.tables
    elif step1 == "14":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts14.docx")
        table1 = doc1.tables
    elif step1 == "0":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts0.docx")
        table1 = doc1.tables
    elif step1 == "1":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts1.docx")
        table1 = doc1.tables
    elif step1 == "15":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts15.docx")
        table1 = doc1.tables
    elif step1 == "16":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts16.docx")
        table1 = doc1.tables
    elif step1 == "17":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts17.docx")
        table1 = doc1.tables
    elif step1 == "18":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts18.docx")
        table1 = doc1.tables
    elif step1 == "19":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts19.docx")
        table1 = doc1.tables
    elif step1 == "20":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts20.docx")
        table1 = doc1.tables


    idTest1 = table1[0].rows[0].cells[2]
    add_idScenario = idTest1.add_paragraph(idTest)
    user_story = table1[0].rows[4].cells[2]
    add_UserStory = user_story.add_paragraph(idTest)

    tabela_template = table1[0].rows

    tabela_template = tabela_template[6:]

    lista_TestStepow = []
    for row in tabela_template:
        lista_TestStepow.append(row.cells[1])


        #Test Stepy
    index = 0
    for element in lista_TestStepow:
        element.add_paragraph(lista_stepow[index])
        index +=1

        #Test Condition
    lista_TestCondition1 =[]
    for row in tabela_template:
        lista_TestCondition1.append(row.cells[2])
    index = 0
    for element in lista_TestCondition1:
        element.add_paragraph(lista_TestCondition[index])
        index +=1

        #Dodawanie test Condition

    lista_ExpResult1 =[]
    for row in tabela_template:
        lista_ExpResult1.append(row.cells[3])

    index = 0
    for element in lista_ExpResult1:
        element.add_paragraph(lista_ExpResult[index])
        index +=1

        #Dodawanie Excepted Result
    doc1.save(file)