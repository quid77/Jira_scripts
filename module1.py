from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import getpass
import time
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import os
import shutil
import difflib
from glob import glob
import re
import win32com.client as win32
from win32com.client import constants
from docx import Document
import docx
import glob
import fnmatch


paths = glob.glob('C:/Users/jpiet/Downloads/DocFiles/**/*.doc', recursive=True)

def save_as_docx(path):
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate ()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)



for path in paths:
    save_as_docx(path)
os.chdir("C:\\Users\\jpiet\\Downloads\\DocFiles")
doc = os.listdir("C:\\Users\\jpiet\\Downloads\\DocFiles")
for item in doc:
    if item.endswith(".doc"):
        os.remove(item)
time.sleep(10)
#Read word document
path = 'C:\\Users\\jpiet\\Downloads\\DocFiles\\Copy*'
files = glob.glob(path)
os.chdir("C:\\Users\\jpiet\\Downloads")
for file in files:

    doc = docx.Document(file)
    table = doc.tables
    titlle = table[0].rows[0].cells[0].text
    print(titlle)
    title_split = titlle.split()
    print(title_split)
    idTest= title_split[0]
    print(idTest)
    step2= table[2].rows[2].cells[0].text
    print(step2)
    if step2 == "Zephyr Teststep:":
        Testy_zefirowe = table[2].rows[2].cells[1]
    else:
        Testy_zefirowe = table[2].rows[3].cells[1]
    tabela_zefir = Testy_zefirowe.tables
    step1 = tabela_zefir[0].rows[-1].cells[0].text

    print(step1)

    tab_zef = tabela_zefir[0].rows
    tab_zef = tab_zef[1:] # obcinanie tytulu test stepow

    lista_stepow = []
    for row in tab_zef:
        lista_stepow.append(row.cells[1].text)

    lista_stepow[0]
    #Test Condition
    lista_TestCondition=[]

    for row in tab_zef:
        lista_TestCondition.append(row.cells[2].text)

    lista_ExpResult=[]
    for row in tab_zef:
        lista_ExpResult.append(row.cells[3].text)
    if step1 == "2":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts.docx")
        table1 = doc1.tables
    elif step1 == "3":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts3.docx")
        table1 = doc1.tables
    elif step1 == "4":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts4.docx")
        table1 = doc1.tables
    elif step1 == "5":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts5.docx")
        table1 = doc1.tables
    elif step1 == "6":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts6.docx")
        table1 = doc1.tables
    elif step1 == "7":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts7.docx")
        table1 = doc1.tables
    elif step1 == "8":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts8.docx")
        table1 = doc1.tables
    elif step1 == "9":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts9.docx")
        table1 = doc1.tables
    elif step1 == "10":
        doc1 = docx.Document("C:\\Users\\jpiet\\Downloads\\SampleTestScripts10.docx")
        table1 = doc1.tables

    idTest1 = table1[0].rows[0].cells[2]
    add_idScenario = idTest1.add_paragraph(idTest)
    user_story = table1[0].rows[4].cells[2]
    add_UserStory = user_story.add_paragraph(idTest)

    tabela_template = table1[0].rows

    tabela_template = tabela_template[6:]

    lista_TestStepow = []
    for row in tabela_template:
        lista_TestStepow.append(row.cells[1])


        #Test Stepy
    index = 0
    for element in lista_TestStepow:
        element.add_paragraph(lista_stepow[index])
        index +=1

        #Test Condition
    lista_TestCondition1 =[]
    for row in tabela_template:
        lista_TestCondition1.append(row.cells[2])
    index = 0
    for element in lista_TestCondition1:
        element.add_paragraph(lista_TestCondition[index])
        index +=1

        #Dodawanie test Condition

    lista_ExpResult1 =[]
    for row in tabela_template:
        lista_ExpResult1.append(row.cells[3])

    index = 0
    for element in lista_ExpResult1:
        element.add_paragraph(lista_ExpResult[index])
        index +=1

        #Dodawanie Excepted Result
    doc1.save(file)